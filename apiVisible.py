import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document

# Função para substituir texto no documento
def substituir_texto(doc, marcador, substituto):
    for paragrafo in doc.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, substituto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_texto(celula, marcador, substituto)

# Função para carregar a planilha e o modelo
def carregar_planilha_e_modelo():
    global df, caminho_modelo
    try:
        # Selecionar a planilha Excel
        caminho_planilha = filedialog.askopenfilename(filetypes=[("Arquivos Excel", "*.xlsx")])
        df = pd.read_excel(caminho_planilha)

        # Selecionar o modelo de documento Word
        caminho_modelo = filedialog.askopenfilename(filetypes=[("Arquivos Word", "*.docx")])

        # Criar um frame para organizar os widgets
        frame_opcoes = tk.Frame(janela)
        frame_opcoes.pack(fill="x", padx=10)

        # Exibir as colunas disponíveis da planilha para o usuário selecionar
        for coluna in df.columns:
            if "Unnamed" in coluna:
                continue  # Ignorar colunas sem nome

            frame_coluna = tk.Frame(frame_opcoes)
            frame_coluna.pack(side="left", padx=10)

            label = tk.Label(frame_coluna, text=f"Coluna {{ {coluna} }}")
            label.pack()

            lista_opcoes = tk.Listbox(frame_coluna, selectmode=tk.SINGLE, height=len(df.columns))
            for coluna_nome in df.columns:
                if "Unnamed" not in coluna_nome:
                    lista_opcoes.insert(tk.END, coluna_nome)
            lista_opcoes.pack()

            # Adicionar rótulo para exibir a seleção atual
            label_selecao = tk.Label(frame_coluna, text="Nenhuma seleção")
            label_selecao.pack()

            # Função que atualiza a seleção visível
            def on_select(evt, lista=lista_opcoes, label_selecao=label_selecao):
                selecionado = lista.get(lista.curselection())
                label_selecao.config(text=f"Selecionado: {selecionado}")

            # Adicionar o evento de seleção ao Listbox
            lista_opcoes.bind("<<ListboxSelect>>", on_select)

            widgets[coluna] = lista_opcoes

        gerar_botao = tk.Button(janela, text="Gerar Documentos", command=gerar_documentos)
        gerar_botao.pack(pady=20)

    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao carregar a planilha ou modelo: {str(e)}")

# Função para gerar os documentos
def gerar_documentos():
    try:
        for index, row in df.iterrows():
            doc = Document(caminho_modelo)

            for marcador, lista_opcoes in widgets.items():
                coluna_escolhida = lista_opcoes.get(tk.ACTIVE)
                if "Unnamed" in coluna_escolhida:
                    continue
                substituir_texto(doc, f'{{{{ {marcador} }}}}', str(row[coluna_escolhida]))

            doc.save(f'TERMO DE DOACAO - {row["nome"]}.docx')

        messagebox.showinfo("Sucesso", "Documentos gerados com sucesso!")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao gerar documentos: {str(e)}")

# Criar a interface gráfica
janela = tk.Tk()
janela.title("Gerador de Documentos Word")

# Configurar para iniciar em tela cheia
janela.attributes('-fullscreen', True)
janela.bind("<F11>", lambda event: janela.attributes('-fullscreen', True))  # Alternar tela cheia
janela.bind("<Escape>", lambda event: janela.attributes('-fullscreen', False))  # Sair da tela cheia

widgets = {}

btn_carregar = tk.Button(janela, text="Selecionar Planilha e Modelo Word", command=carregar_planilha_e_modelo)
btn_carregar.pack(pady=20)

janela.mainloop()