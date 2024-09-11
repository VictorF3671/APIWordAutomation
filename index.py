import pandas as pd
from docx import Document

# Função para substituir texto no documento
def substituir_texto(doc, marcador, substituto):
    for paragrafo in doc.paragraphs:
        if marcador in paragrafo.text:
            paragrafo.text = paragrafo.text.replace(marcador, substituto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_texto(celula, marcador, substituto)

# Carregar dados da planilha Excel
df = pd.read_excel('lista.xlsx')

# Iterar sobre cada linha da planilha e gerar um documento Word para cada uma
for index, row in df.iterrows():
    # Carregar o documento modelo
    doc = Document('modelo.docx')
    
    # Substituir os marcadores de posição pelo valor correspondente
    substituir_texto(doc, '{{ nome }}', row['nome'])
    substituir_texto(doc, '{{ cpf }}', row['cpf'])
    substituir_texto(doc, '{{ rg }}', row['rg'])
    substituir_texto(doc, '{{ endereco }}', row['endereço'])
    substituir_texto(doc, '{{ bairro }}', row['bairro'])
    substituir_texto(doc, '{{ cep }}', row['cep'])
    substituir_texto(doc, '{{ telefone }}', row['telefone'])   
    # Salvar o documento resultante
    doc.save(f'TERMO DE DOACAO - {row['nome']}.docx')

print("Documentos gerados com sucesso!")